import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = Document()

    # Title
    title = doc.add_heading('NOK Inc: WordPress to Next.js Migration Process', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('This document explains the complete process of migrating the NOK Inc website from WordPress to a modern Next.js architecture. The guide is written in simple terms, explaining the technical concepts used along the way.')

    # Section 1: Introduction & Terminologies
    doc.add_heading('1. Introduction and Terminologies', level=1)
    
    doc.add_heading('Why we migrated:', level=2)
    doc.add_paragraph('WordPress is great for basic websites, but as NOK Inc grows, it requires a faster, more secure, and highly customizable platform. We moved to a modern web stack to ensure the website loads instantly, looks fantastic on all devices, and is easier to maintain in the long run.')

    doc.add_heading('Key Technical Terms (Explained Simply):', level=2)
    terms = [
        ('Next.js', 'A powerful framework (a set of tools) built on top of React. It allows us to build extremely fast websites by preparing the web pages in advance before a user even visits them.'),
        ('React', 'A technology created by Facebook used to build user interfaces. It lets us break the website into small, reusable building blocks called "components" (like a Lego set).'),
        ('Tailwind CSS', 'A styling tool that lets us design the website directly within our code. Instead of writing separate styling files, we use simple descriptive words (like "text-center" or "bg-blue-500") to make the website look beautiful.'),
        ('Component', 'A reusable piece of the website. For example, the Navigation Bar at the top and the Footer at the bottom are components. We build them once and reuse them on every page.'),
        ('Static Export', 'A process where we turn the entire website into simple, lightweight files that can be hosted anywhere cheaply and securely, making the site incredibly fast and immune to most hacking attempts.')
    ]
    
    for term, definition in terms:
        p = doc.add_paragraph()
        p.add_run(f'{term}: ').bold = True
        p.add_run(definition)

    # Section 2: Step-by-Step Process
    doc.add_heading('2. The Migration Process (Step-by-Step)', level=1)

    # Step 1
    doc.add_heading('Step 1: Setting up the Next.js Foundation', level=2)
    doc.add_paragraph('We started by creating a brand-new Next.js project. We configured the system to use "Static Export," which means the final website will be extremely fast and secure. We also cleaned up the default code that comes with Next.js to start with a blank canvas.')

    # Step 2
    doc.add_heading('Step 2: Rescuing Assets from WordPress', level=2)
    doc.add_paragraph('A website needs its images! We wrote a special script that automatically scanned the old WordPress website, found all the important images (like the NOK Inc logo, background images, and product photos), and safely downloaded them into our new project folder so we could reuse them.')

    # Step 3
    doc.add_heading('Step 3: Building the Global Layout (The Skeleton)', level=2)
    doc.add_paragraph('Every page on a website usually shares the same Header (the menu at the top) and Footer (the links at the bottom). We built a "Global Layout". This acts like a picture frame; no matter what page you navigate to, the Header and Footer stay perfectly in place while the content in the middle changes.')

    # Step 4
    doc.add_heading('Step 4: Designing Reusable Components (The Lego Blocks)', level=2)
    doc.add_paragraph('Instead of rebuilding the same things over and over, we created reusable building blocks:')
    doc.add_paragraph('- The Header: We made it "fixed" so it floats beautifully at the top of the screen even when you scroll. We added a "glassmorphism" effect (making it slightly transparent and blurry like frosted glass) to give it a premium feel.', style='List Bullet')
    doc.add_paragraph('- The Footer: We organized the contact information, social media links, and quick links into neat columns.', style='List Bullet')
    doc.add_paragraph('- The Section: A master container that ensures all text and images are perfectly aligned in the center of the screen with proper spacing on all sides.', style='List Bullet')

    # Step 5
    doc.add_heading('Step 5: Rebuilding the Pages', level=2)
    doc.add_paragraph('We meticulously rebuilt each page using our new Lego blocks and Tailwind CSS for styling:')
    doc.add_paragraph('- Home Page: We added a massive, edge-to-edge "Hero" image at the top with a dark gradient overlay so the white text stands out perfectly. We also added a sleek "Contact Us" banner at the bottom with curved edges and interactive buttons.', style='List Bullet')
    doc.add_paragraph('- Products & Services: We arranged the offerings into neat, uniform cards. Each card has an image, a title, and bullet points explaining the benefits.', style='List Bullet')
    doc.add_paragraph('- Contact Page: We built a clean layout displaying your office address, phone numbers, and a professional message form.', style='List Bullet')

    # Step 6
    doc.add_heading('Step 6: Quality Assurance and Bug Fixing', level=2)
    doc.add_paragraph('Finally, we ran automated tools to check our work for any mistakes:')
    doc.add_paragraph('- Linter (ESLint): This acts like a grammar checker for code. It found a couple of minor issues (like an unused image tool and a broken image link), which we immediately fixed.', style='List Bullet')
    doc.add_paragraph('- Type Checker (TypeScript): This ensures that all the code speaks the same language and fits together perfectly without errors. The website passed this check with flying colors.', style='List Bullet')

    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph('The NOK Inc website is now fully migrated to a state-of-the-art Next.js framework. It is faster, more secure, and beautifully styled with Tailwind CSS, ready to scale alongside your business.')

    # Save to Desktop
    desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop', 'NOK_Inc_Migration_Process.docx')
    doc.save(desktop_path)
    print(f"Document saved successfully to: {desktop_path}")

if __name__ == "__main__":
    create_document()
